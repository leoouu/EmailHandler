import os
import fitz  # PyMuPDF
import docx
import re
import database

def process_file(file_path, keywords):
    if file_path.endswith('.pdf'):
        return process_pdf(file_path, keywords)
    elif file_path.endswith('.docx'):
        return process_docx(file_path, keywords)
    else:
        return "Formato não suportado."

def process_pdf(file_path, keywords):
    with fitz.open(file_path) as doc:
        text = ""
        for page in doc:
            text += page.get_text()
    candidate_name = extract_candidate_name(text)
    email = extract_email(text)
    database.save_resume(candidate_name, file_path, email, text, "manual")
    return f"{candidate_name} - {file_path}"

def process_docx(file_path, keywords):
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    candidate_name = extract_candidate_name(text)
    email = extract_email(text)
    database.save_resume(candidate_name, file_path, email, text, "manual")
    return f"{candidate_name} - {file_path}"

def extract_candidate_name(text):
    lines = text.split('\n')
    for line in lines:
        if re.match(r'^[A-Z][a-z]*\s[A-Z][a-z]*', line):
            return line.strip()
    return "Nome não encontrado"

def extract_email(text):
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    return email_match.group(0) if email_match else "Email não encontrado"

def read_file(file_path):
    if file_path.endswith('.pdf'):
        return read_pdf(file_path)
    elif file_path.endswith('.docx'):
        return read_docx(file_path)
    else:
        return "Formato não suportado."

def read_pdf(file_path):
    with fitz.open(file_path) as doc:
        text = ""
        for page in doc:
            text += page.get_text()
    return text

def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])
